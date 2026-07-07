import os
from pathlib import Path

os.environ.setdefault("KMP_DUPLICATE_LIB_OK", "TRUE")

import pandas as pd
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
RESULT_DIR = ROOT / "results" / "pilot_physionet"
INPUT_DOCX = ROOT / "代理脑论文_补充实验与整体框架图_完善版.docx"
OUTPUT_DOCX = ROOT / "代理脑论文_补充实验与整体框架图_实验结果版.docx"


def fmt_pct(x):
    return f"{100 * float(x):.2f}"


def fmt_num(x):
    if pd.isna(x):
        return "N/A"
    return f"{float(x):.4f}"


def replace_main_table(table, metrics):
    rows = [
        ["PhysioNet pilot LOSO", "Bandpower-LDA"],
        ["PhysioNet pilot LOSO", "Raw-LogReg"],
        ["PhysioNet pilot LOSO", "EEGNet-Lite"],
        ["PhysioNet pilot LOSO", "HAB-Prototype"],
    ]
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    for protocol, method in rows:
        m = metrics.loc[metrics["method"] == method].iloc[0]
        cells = table.add_row().cells
        vals = [
            protocol,
            method,
            fmt_pct(m["acc"]),
            fmt_num(m["macro_f1"]),
            fmt_num(m["kappa"]),
            fmt_num(m["nll"]),
            fmt_num(m["ece"]),
            fmt_num(m["brier"]),
            fmt_num(m["params_m"]),
            fmt_num(m["latency_ms"]),
        ]
        for cell, val in zip(cells, vals):
            cell.text = val


def replace_world_table(table, wm):
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    for _, row in wm.iterrows():
        cells = table.add_row().cells
        vals = [
            row["method"],
            fmt_num(row["representation_mse"]),
            "N/A",
            "N/A",
            row["note"],
        ]
        for cell, val in zip(cells, vals):
            cell.text = val


def replace_risk_table(table, risk, ece):
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    for _, row in risk.iterrows():
        cells = table.add_row().cells
        vals = [
            "HAB-Prototype",
            f"{100 * row['target_coverage']:.0f}%",
            f"{100 * row['actual_coverage']:.2f}%",
            fmt_num(row["selective_acc"]),
            fmt_num(row["selective_risk"]),
            fmt_num(row["expected_cost"]),
            f"{100 * row['confirmation_rate']:.2f}%",
            fmt_num(ece),
        ]
        for cell, val in zip(cells, vals):
            cell.text = val


def add_pilot_note(doc):
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "5. Experiments":
            note = doc.paragraphs[i + 1].insert_paragraph_before(
                "说明：下列表格已填入一组可复现的 pilot 实验结果。实验范围为 PhysioNet EEGBCI 数据集前 5 名受试者、左/右手运动想象二分类、留第 5 名受试者测试。该结果用于验证代码流水线与论文表格格式，不应替代 BNCI2014_001 全量主实验结论。"
            )
            return note


def main():
    metrics = pd.read_csv(RESULT_DIR / "main_metrics.csv")
    risk = pd.read_csv(RESULT_DIR / "risk_metrics.csv")
    wm = pd.read_csv(RESULT_DIR / "world_model_metrics.csv")
    doc = Document(INPUT_DOCX)
    add_pilot_note(doc)
    replace_main_table(doc.tables[3], metrics)
    replace_world_table(doc.tables[4], wm)
    hab_ece = float(metrics.loc[metrics["method"] == "HAB-Prototype", "ece"].iloc[0])
    replace_risk_table(doc.tables[7], risk, hab_ece)
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX.resolve())


if __name__ == "__main__":
    main()
