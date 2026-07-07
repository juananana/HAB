from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib import font_manager
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Circle
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(".")
SOURCE_DOCX = max(
    [
        p
        for p in ROOT.glob("*.docx")
        if not p.name.startswith("~$") and p.stat().st_size > 100000 and "完善版" not in p.name
    ],
    key=lambda p: p.stat().st_size,
)
OUTPUT_DOCX = ROOT / "代理脑论文_补充实验与整体框架图_完善版.docx"
FIGURE_PATH = ROOT / "hab_overall_framework_final.png"


def setup_font():
    candidates = [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            font_manager.fontManager.addfont(candidate)
            plt.rcParams["font.family"] = font_manager.FontProperties(fname=candidate).get_name()
            break
    plt.rcParams["axes.unicode_minus"] = False


def box(ax, xy, wh, text, fc, ec, fontsize=10, lw=1.4, radius=0.04):
    x, y = xy
    w, h = wh
    patch = FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle=f"round,pad=0.018,rounding_size={radius}",
        facecolor=fc,
        edgecolor=ec,
        linewidth=lw,
    )
    ax.add_patch(patch)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=fontsize, color="#182033")
    return patch


def blank_box(ax, xy, wh, fc, ec, lw=1.4, radius=0.04):
    x, y = xy
    w, h = wh
    patch = FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle=f"round,pad=0.018,rounding_size={radius}",
        facecolor=fc,
        edgecolor=ec,
        linewidth=lw,
    )
    ax.add_patch(patch)
    return patch


def arrow(ax, start, end, color="#3d4a5c", style="-|>", lw=1.8, dashed=False, rad=0):
    arr = FancyArrowPatch(
        start,
        end,
        arrowstyle=style,
        mutation_scale=16,
        linewidth=lw,
        color=color,
        linestyle=(0, (4, 3)) if dashed else "solid",
        connectionstyle=f"arc3,rad={rad}",
    )
    ax.add_patch(arr)


def draw_graph(ax, cx, cy, scale=1.0):
    points = [
        (cx - 0.035 * scale, cy + 0.035 * scale),
        (cx + 0.015 * scale, cy + 0.055 * scale),
        (cx + 0.055 * scale, cy + 0.010 * scale),
        (cx + 0.015 * scale, cy - 0.050 * scale),
        (cx - 0.055 * scale, cy - 0.020 * scale),
        (cx - 0.010 * scale, cy),
    ]
    edges = [(0, 1), (1, 2), (2, 3), (3, 4), (4, 0), (0, 5), (1, 5), (2, 5), (3, 5)]
    for i, j in edges:
        ax.plot([points[i][0], points[j][0]], [points[i][1], points[j][1]], color="#8a96a8", lw=1)
    for x, y in points:
        ax.add_patch(Circle((x, y), 0.010 * scale, facecolor="#1f77b4", edgecolor="white", lw=0.8))


def make_framework_figure():
    setup_font()
    fig, ax = plt.subplots(figsize=(14, 8), dpi=220)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    ax.text(0.5, 0.955, "Hierarchical Agentic Brain 总体框架", ha="center", va="center",
            fontsize=22, fontweight="bold", color="#101a37")
    ax.text(0.5, 0.915, "从多模态脑信号到认知状态预测、风险校准与闭环交互决策", ha="center",
            va="center", fontsize=12, color="#59616f")

    box(ax, (0.035, 0.66), (0.17, 0.13), "外部上下文\n任务环境 $E_t$\n历史动作/反馈 $A_{t-1}$", "#f5f9ff", "#2f66d0", 10)
    box(ax, (0.035, 0.43), (0.17, 0.15), "脑信号输入\nEEG / fMRI / fNIRS\n窗口 $X_{t-w:t}$", "#f5f9ff", "#2f66d0", 10)
    box(ax, (0.035, 0.23), (0.17, 0.11), "训练与验证协议\n数据划分 / 预处理\n指标与安全阈值", "#f8f8fb", "#9aa3b2", 9)

    box(ax, (0.28, 0.18), (0.43, 0.64), "", "#fbfdff", "#16808b", lw=1.8, radius=0.035)
    ax.text(0.495, 0.795, "层级认知建模核心", ha="center", va="center", fontsize=13,
            fontweight="bold", color="#0c5b63")

    box(ax, (0.315, 0.59), (0.36, 0.14),
        "3  小世界认知世界模型\n预测 $\\widehat{\\mathcal{G}}^{C}_{t+1}=T(\\mathcal{G}^{C}_{t},A_{t},E_{t})$\n约束局部聚类、全局整合与稀疏结构",
        "#fff7e8", "#e7a83e", 9.3)
    draw_graph(ax, 0.35, 0.665, 0.9)

    blank_box(ax, (0.315, 0.39), (0.36, 0.15), "#eefaf1", "#5aa96b")
    ax.text(0.405, 0.49, "2  神经-认知单元抽象", ha="center", va="center", fontsize=10, color="#182033")
    ax.text(0.405, 0.455, "$C_t=\\phi(\\mathcal{G}^{N}_{t})$", ha="center", va="center", fontsize=10, color="#182033")
    ax.text(0.405, 0.42, "潜变量：注意、记忆、规则、意图、情绪、不确定性",
            ha="center", va="center", fontsize=8.4, color="#182033")
    labels = ["注意", "记忆", "规则", "意图", "情绪", "不确定性"]
    coords = [(0.545, 0.505), (0.610, 0.505), (0.645, 0.465), (0.610, 0.425), (0.545, 0.425), (0.515, 0.465)]
    for (x, y), label in zip(coords, labels):
        ax.add_patch(Circle((x, y), 0.024, facecolor="#e9f2ff", edgecolor="#598bd4", lw=1.2))
        ax.text(x, y, label, ha="center", va="center", fontsize=7.2)
    for i in range(len(coords)):
        for j in range(i + 1, len(coords)):
            if (i + j) % 2 == 0:
                ax.plot([coords[i][0], coords[j][0]], [coords[i][1], coords[j][1]], color="#a3adba", lw=0.7, alpha=0.75)

    box(ax, (0.315, 0.22), (0.36, 0.12),
        "1  信号预处理与动态神经图\n滤波、标准化、窗口化、动态邻接学习\n$\\mathcal{G}^{N}_{t}=(V^{N}_{t},A^{N}_{t},H^{N}_{t})$",
        "#edf5ff", "#5d8ed8", 9.3)

    box(ax, (0.78, 0.60), (0.17, 0.16), "任务与意图解码\n$\\hat{y}_t=f_y(\\mathcal{G}^{C}_{t})$\nAccuracy / F1 / Kappa", "#fff6f6", "#d14b45", 10)
    box(ax, (0.78, 0.39), (0.17, 0.15), "不确定性校准\n$u_t=f_u(\\mathcal{G}^{C}_{t})$\nNLL / ECE / Brier", "#fff6f6", "#d14b45", 10)
    box(ax, (0.78, 0.19), (0.17, 0.13), "代理交互动作\n$a_t=\\pi(\\hat{y}_t,u_t,E_t)$\n执行 / 确认 / 延迟", "#fff6f6", "#d14b45", 10)

    box(ax, (0.16, 0.055), (0.68, 0.075),
        "联合训练目标：$\\mathcal{L}=\\mathcal{L}_{task}+\\lambda_1\\mathcal{L}_{wm}+\\lambda_2\\mathcal{L}_{sw}+\\lambda_3\\mathcal{L}_{unc}+\\lambda_4\\mathcal{L}_{risk}+\\lambda_5\\mathcal{L}_{align}$",
        "#fbf8ff", "#7f65c5", 10.5)

    arrow(ax, (0.205, 0.72), (0.28, 0.66))
    arrow(ax, (0.205, 0.50), (0.28, 0.29))
    arrow(ax, (0.495, 0.34), (0.495, 0.39), color="#4d85d3")
    arrow(ax, (0.495, 0.54), (0.495, 0.59), color="#4fa763")
    arrow(ax, (0.71, 0.66), (0.78, 0.68))
    arrow(ax, (0.71, 0.48), (0.78, 0.465))
    arrow(ax, (0.865, 0.60), (0.865, 0.54), color="#d14b45")
    arrow(ax, (0.865, 0.39), (0.865, 0.32), color="#d14b45")
    arrow(ax, (0.78, 0.25), (0.72, 0.30), color="#5b6573", dashed=True, rad=0)

    plt.savefig(FIGURE_PATH, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def delete_from_paragraph(doc, start_idx):
    body = doc._body._element
    for element in list(body)[start_idx:]:
        if element.tag.endswith("sectPr"):
            continue
        body.remove(element)


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    return p


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.italic = True
        run.font.size = Pt(10)
    return p


def add_table(doc, caption, headers, rows):
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()
    return table


def rebuild_doc():
    make_framework_figure()
    doc = Document(SOURCE_DOCX)

    start = None
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "4. Methodology":
            start = i
            break
    if start is None:
        raise RuntimeError("Could not find section 4.")
    delete_from_paragraph(doc, start)

    add_heading(doc, "4. Methodology", 1)
    add_heading(doc, "4.1 Overview of Hierarchical Agentic Brain", 2)
    add_para(doc, "如图 1 所示，Hierarchical Agentic Brain（HAB）由输入与上下文建模、动态神经图编码、神经—认知单元抽象、小世界认知世界模型、不确定性感知代理决策以及联合训练目标六个部分组成。该框架的核心思想不是直接把脑信号窗口映射为类别标签，而是先将低层神经活动组织为动态结构，再抽象为可追踪的认知单元图，并在此基础上完成意图解码、状态预测、风险评估和交互动作选择。")
    add_para(doc, "与传统 BCI 的一次性分类器相比，HAB 明确引入历史动作、外部任务环境和反馈信息，使系统能够形成“当前观测—潜在认知状态—下一状态预测—交互决策—反馈更新”的闭环链路。该设计与第 3 节中的多目标优化问题保持一致：任务预测损失负责保证解码性能，世界模型预测损失刻画认知状态演化，小世界结构约束提升认知图的局部聚合与全局整合能力，不确定性与风险损失用于支持更可信的代理决策。")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIGURE_PATH), width=Inches(6.2))
    add_caption(doc, "图 1  Hierarchical Agentic Brain 总体框架")

    add_heading(doc, "4.2 Dynamic Neural Graph Encoder", 2)
    add_para(doc, "动态神经图编码器以滑动时间窗口内的脑信号 X_{t-w:t} 为输入。对 EEG 场景而言，节点通常对应电极通道；对 fMRI 或 fNIRS 场景，节点可对应脑区或功能通道。模型首先通过时序卷积、频带滤波或轻量级 Transformer 提取每个节点的局部时频特征，得到 H_t^N。随后根据特征相似度、可学习注意力或稀疏邻居选择构造动态邻接矩阵 A_t^N，从而形成神经图 G_t^N=(V_t^N,A_t^N,H_t^N)。")
    add_para(doc, "为了避免动态图退化为完全连接图，实际实现中可限制每个节点仅保留 top-k 条边，或对邻接矩阵加入稀疏正则。动态图编码层输出的图级表征不仅用于任务分类，也作为后续认知单元抽象和神经—认知对齐损失的输入。")

    add_heading(doc, "4.3 Neural-to-Cognitive Unit Abstraction", 2)
    add_para(doc, "神经—认知单元抽象模块将动态神经图映射为 K 个潜在认知单元。每个认知单元由嵌入向量、激活强度、不确定性估计和结构角色特征组成，可被理解为模型学习到的中间认知变量。该模块可以采用跨注意力机制实现：以一组可学习的认知查询向量作为 query，以神经图节点表征作为 key 和 value，从而得到 C_t=[c_1^t,...,c_K^t]。")
    add_para(doc, "需要强调的是，认知单元是可解释的潜变量，而不是直接观测到的心理学标签。因此，论文中应避免将某个单元直接等同于真实的注意、记忆或情绪机制。更严谨的表述是：若某些单元在不同任务类别、时间阶段或不确定性水平下呈现稳定差异，则说明它们与相应认知状态存在统计关联。")

    add_heading(doc, "4.4 Small-world Cognitive World Model", 2)
    add_para(doc, "小世界认知世界模型在认知单元图 G_t^C 上建模状态转移。该模块一方面根据当前认知图、历史动作 A_t 和外部上下文 E_t 预测下一时刻认知图；另一方面通过小世界结构约束鼓励认知图同时具备较高聚类系数、较短平均路径长度、适度模块化和稀疏连接。")
    add_para(doc, "在实验中，世界模型的有效性不能只通过最终分类准确率间接说明，还应单独报告下一状态预测误差、下一窗口重构误差或基于预测状态的下一步分类性能。这样可以证明该模块确实学习到状态演化规律，而不是仅作为普通正则项存在。")

    add_heading(doc, "4.5 Uncertainty-Aware Agentic Decision Module", 2)
    add_para(doc, "不确定性感知代理决策模块输出三个量：任务或意图预测 \\hat{y}_t、预测不确定性 u_t 以及交互动作 a_t。当不确定性低于阈值时，系统执行预测控制指令；当不确定性较高时，系统请求用户确认或延迟执行。阈值必须仅在验证集上选择，测试集只用于一次性报告，以避免高估安全决策能力。")
    add_para(doc, "对于离线公开数据，本文只能进行风险决策模拟，不能直接宣称完成真实闭环 BCI 验证。因此，实验章节应将结论限定为“离线条件下模型具备风险感知与选择性执行能力”。若后续需要强调在线交互能力，应增加伪在线滑动窗口评估或小规模用户闭环实验。")

    add_heading(doc, "4.6 Training Objective and Optimization", 2)
    add_para(doc, "训练阶段采用端到端联合优化。总损失由任务预测、认知世界模型预测、小世界结构、不确定性校准、交互风险和神经—认知对齐损失组成。为了保证各项损失可比，建议先单独训练任务解码主干，再逐步加入世界模型和结构约束，并在验证集上搜索 λ_1 至 λ_5 的权重。")
    add_para(doc, "正式实验需要记录完整训练配置，包括随机种子、数据划分、滤波范围、窗口长度、学习率、batch size、早停规则、最佳 checkpoint 选择规则、参数量统计方法和 batch size=1 的推理时延。安全阈值、校准温度和代价矩阵均应由验证集确定。")

    add_heading(doc, "5. Experiments", 1)
    add_para(doc, "本节给出用于验证 HAB 框架的实验方案。由于当前稿件尚未包含真实代码运行日志和性能数值，本文不写入未经验证的结果，而是补全可执行的实验协议、对比模型、评价指标、消融路径和结果表结构。正式实验完成后，应使用真实运行结果替换表格中的“待填”，并在附录或代码仓库中固定最终配置。")
    add_para(doc, "实验围绕四个问题展开：第一，层级化代理脑是否提升基础脑信号解码性能；第二，认知世界模型是否能够预测状态演化；第三，小世界约束是否改善认知图结构质量与泛化能力；第四，不确定性感知策略是否能够在降低高风险错误执行的同时维持可接受的交互覆盖率。")

    add_heading(doc, "5.1 Datasets and Experimental Settings", 2)
    add_para(doc, "首轮实验建议以 EEG 运动想象任务作为主验证场景。该任务具有公开数据、明确类别标签和成熟基线，适合检验层级建模、小世界结构约束和安全决策模块是否带来可复核收益。框架虽然支持 fMRI、fNIRS 和多模态输入，但在缺少统一实验协议前，不应把多模态能力写成已经验证的结论。")
    add_table(
        doc,
        "表 1  数据集与实验用途",
        ["数据集", "信号形式与规模", "任务定义", "本文中的用途"],
        [
            ["BNCI2014_001 / BCI Competition IV 2a", "9 名受试者；2 次独立会话；22 路 EEG + 3 路 EOG；250 Hz", "左手、右手、双脚、舌部运动想象四分类", "主实验数据集；用于被试内跨会话、跨被试留一、消融和可解释性分析"],
            ["PhysioNet EEGMMIDB", "109 名志愿者；64 路 EEG；每名受试者多个运动/想象 run", "真实运动与运动想象任务；需整理为兼容子任务", "外部鲁棒性验证；检验不同采集系统和通道数量下的稳定性"],
            ["SEED（可选扩展）", "15 名受试者；62 路 EEG；影片诱发情绪数据", "正向、中性、负向情绪三分类", "跨任务扩展验证；不作为首轮核心结论"],
        ],
    )
    add_para(doc, "BNCI2014_001 的两次会话在不同日期采集，适合构造被试内跨会话评估：第一会话用于训练与验证，第二会话只用于最终测试。跨被试实验采用留一被试法，即每次选择一名受试者作为测试对象，其余受试者用于训练，并在训练集合内部划分验证集。PhysioNet EEGMMIDB 应根据任务范式重新整理兼容子任务，避免把真实运动、运动想象和静息段混入同一结果表。")
    add_table(
        doc,
        "表 2  需要固化的实验配置",
        ["配置项", "建议设置", "正式实验需确认"],
        [
            ["数据划分", "被试内跨会话 + 跨被试留一法", "训练、验证、测试清单与随机划分文件"],
            ["预处理", "保留 EEG 通道；带通滤波；逐通道标准化；固定窗口切分", "滤波范围、窗口起止时间、步长、重叠率、是否重采样"],
            ["训练策略", "统一优化器、早停和验证集选模", "优化器、学习率、batch size、最大轮数、patience"],
            ["重复实验", "不少于 5 个随机种子", "种子列表、均值、标准差和显著性检验"],
            ["运行环境", "报告参数量与单样本推理时延", "CPU/GPU 型号、软件版本、batch size=1 推理脚本"],
        ],
    )

    add_heading(doc, "5.2 Baselines and Evaluation Metrics", 2)
    add_para(doc, "对比方法需要覆盖传统特征工程、端到端深度网络、图结构模型和本文完整模型四类路线。所有方法使用相同的数据划分、预处理输入和验证集调参规则。传统方法用于给出经典 BCI 参照，卷积网络用于比较主流端到端解码器，图结构模型用于分离动态神经图和认知世界模型的贡献。")
    add_table(
        doc,
        "表 3  对比方法设置",
        ["类别", "方法", "对比目的"],
        [
            ["传统方法", "CSP + LDA", "检验模型相对于经典空间滤波方法的收益"],
            ["传统方法", "FBCSP + LDA", "提供多频带空间特征基线"],
            ["卷积网络", "EEGNet", "提供轻量级端到端 EEG 解码基线"],
            ["卷积网络", "ShallowConvNet", "比较偏向频带功率特征的浅层卷积模型"],
            ["卷积网络", "DeepConvNet", "比较更深的时空特征提取模型"],
            ["图结构模型", "Dynamic Neural Graph Only", "保留动态神经图编码，移除认知单元和世界模型"],
            ["本文模型", "Hierarchical Agentic Brain", "完整模型，包含动态图、认知单元、小世界世界模型、校准和风险策略"],
        ],
    )
    add_para(doc, "主要指标包括 Accuracy、Macro-F1、Cohen's kappa、NLL、ECE、Brier score、参数量和单样本推理时延。安全决策实验还需报告 Coverage、Selective Accuracy、Selective Risk、Expected Cost 和 Confirmation Rate。对于跨受试者实验，应同时报告均值、标准差和每名受试者的明细结果。")

    add_heading(doc, "5.3 Main Comparison Experiments", 2)
    add_para(doc, "主对比实验首先验证 HAB 的任务解码能力。被试内跨会话实验用于评估模型对同一用户跨天信号漂移的适应性；跨被试留一实验用于评估泛化到新用户时的稳定性。若最终结果显示 HAB 在准确率上提升有限，也应结合 NLL、ECE 和风险指标分析其是否在可信交互方面具有优势。")
    add_table(
        doc,
        "表 4  主实验结果模板（真实运行后填写）",
        ["评估协议", "方法", "Acc.(%)", "Macro-F1", "Kappa", "NLL↓", "ECE↓", "Brier↓", "参数量(M)", "时延(ms)"],
        [
            ["被试内跨会话", "CSP-LDA", "待填", "待填", "待填", "待填", "待填", "待填", "待填", "待填"],
            ["被试内跨会话", "FBCSP-LDA", "待填", "待填", "待填", "待填", "待填", "待填", "待填", "待填"],
            ["被试内跨会话", "EEGNet", "待填", "待填", "待填", "待填", "待填", "待填", "待填", "待填"],
            ["被试内跨会话", "DNG Only", "待填", "待填", "待填", "待填", "待填", "待填", "待填", "待填"],
            ["被试内跨会话", "HAB (Ours)", "待填", "待填", "待填", "待填", "待填", "待填", "待填", "待填"],
            ["跨被试留一法", "EEGNet", "待填", "待填", "待填", "待填", "待填", "待填", "待填", "待填"],
            ["跨被试留一法", "DNG Only", "待填", "待填", "待填", "待填", "待填", "待填", "待填", "待填"],
            ["跨被试留一法", "HAB (Ours)", "待填", "待填", "待填", "待填", "待填", "待填", "待填", "待填"],
        ],
    )
    add_para(doc, "为了验证认知世界模型并非仅作为附加正则项存在，需要单独评估下一时刻状态预测。若真实认知状态不可直接观测，可采用下一窗口神经图表征误差、下一窗口脑信号重构误差或基于预测状态完成下一步分类的准确率作为代理指标。")
    add_table(
        doc,
        "表 5  认知世界模型预测结果模板",
        ["方法", "表征预测误差↓", "重构误差↓", "下一步 Acc.↑", "说明"],
        [
            ["Last-state Copy", "待填", "待填", "待填", "直接复制当前状态作为下一步预测"],
            ["GRU Predictor", "待填", "待填", "待填", "使用序列模型预测下一窗口状态"],
            ["HAB (Ours)", "待填", "待填", "待填", "使用认知单元图、历史动作和环境信息预测状态演化"],
        ],
    )

    add_heading(doc, "5.4 Ablation Study", 2)
    add_para(doc, "消融实验用于定位性能和可信性收益来自哪些模块。除删除或替换指定组件外，其余训练设置、随机种子和数据划分保持一致。对于小世界结构、不确定性校准和风险策略，不能只观察分类准确率，还应结合结构指标和选择性风险指标判断模块是否真正发挥作用。")
    add_table(
        doc,
        "表 6  消融实验模板",
        ["模型变体", "删除或替换内容", "Acc.", "Macro-F1", "ECE↓", "Selective Risk↓", "Small-worldness σ↑"],
        [
            ["完整模型", "保留全部模块", "待填", "待填", "待填", "待填", "待填"],
            ["w/o Neural Graph", "将动态图编码替换为通道独立时序编码", "待填", "待填", "待填", "待填", "待填"],
            ["w/o Cognitive Units", "直接从神经图池化表征完成分类", "待填", "待填", "待填", "待填", "待填"],
            ["w/o SW Constraint", "移除小世界结构正则项", "待填", "待填", "待填", "待填", "待填"],
            ["w/o WM Prediction", "移除下一状态预测损失", "待填", "待填", "待填", "待填", "待填"],
            ["w/o Calibration", "保留分类概率但不做温度缩放或校准损失", "待填", "待填", "待填", "待填", "待填"],
            ["w/o Risk Policy", "始终执行最大概率类别", "待填", "待填", "待填", "待填", "待填"],
        ],
    )

    add_heading(doc, "5.5 Optimization and Hyperparameter Analysis", 2)
    add_para(doc, "超参数分析应围绕论文主张展开，而不是进行无目的的大规模穷举。建议先在验证集上进行粗粒度搜索，再对敏感区间局部细化。每组设置至少重复多个随机种子，并报告均值和标准差。")
    add_table(
        doc,
        "表 7  关键超参数分析",
        ["超参数", "候选范围", "分析目的"],
        [
            ["认知单元数量 K", "4 / 8 / 12 / 16", "检验认知抽象粒度与性能、可解释性的平衡"],
            ["动态图邻居数 k", "4 / 8 / 12", "观察过稀或过密连接对神经图建模的影响"],
            ["小世界损失权重 λ_sw", "0 / 1e-3 / 1e-2 / 1e-1", "检验结构先验强度，避免图结构退化"],
            ["世界模型损失权重 λ_wm", "0 / 1e-3 / 1e-2 / 1e-1", "检验时序预测约束是否改善泛化"],
            ["不确定性阈值 τ", "验证集分位点或目标覆盖率", "控制执行、确认和延迟之间的权衡"],
            ["时间窗口与步长", "由采样率和任务阶段确定", "观察短时响应与长时状态平滑之间的差异"],
        ],
    )
    add_para(doc, "训练曲线至少应包含任务损失、世界模型预测损失、验证集准确率、NLL 和 ECE。若加入小世界约束后准确率下降但结构指标改善，需要检查正则权重是否过大；若认知图接近完全连接，则应提高稀疏约束或限制邻居数。")

    add_heading(doc, "5.6 Uncertainty and Safe Decision Analysis", 2)
    add_para(doc, "公开离线数据无法直接证明真实用户参与下的闭环交互效果，但可以构造可复现的风险决策模拟。对每个测试窗口，模型输出类别概率和不确定性；当不确定性低于阈值 τ 时执行预测动作，当不确定性较高时请求用户确认或延迟决策。逐步改变 τ 可获得覆盖率—风险曲线，并观察模型是否将易错样本优先交给确认机制处理。")
    add_table(
        doc,
        "表 8  风险感知决策实验模板",
        ["策略", "目标覆盖率", "实际覆盖率", "Selective Acc.", "Selective Risk↓", "Expected Cost↓", "确认率", "ECE↓"],
        [
            ["Always Execute", "100%", "待填", "待填", "待填", "待填", "0%", "待填"],
            ["Entropy-τ", "90%", "待填", "待填", "待填", "待填", "待填", "待填"],
            ["Entropy-τ", "80%", "待填", "待填", "待填", "待填", "待填", "待填"],
            ["Calibrated-τ", "90%", "待填", "待填", "待填", "待填", "待填", "待填"],
            ["Calibrated-τ", "80%", "待填", "待填", "待填", "待填", "待填", "待填"],
        ],
    )
    add_para(doc, "风险代价矩阵不应随意固定为单一数值。建议根据应用场景设置多组代价比例，分别考察“错误执行代价显著高于确认代价”和“需要维持较高响应覆盖率”两类条件。上述模拟只能支持离线风险感知证据，不能替代真实在线 BCI 用户实验。")

    add_heading(doc, "5.7 Interpretability of Cognitive Units and Small-World Structure", 2)
    add_para(doc, "可解释性分析分为三个层次：神经图层面展示动态连接矩阵和关键通道拓扑，认知单元层面展示不同类别下的激活热力图、连接强度和时间轨迹，决策层面分析高不确定性样本对应的认知图状态，并比较执行、确认和延迟决策时的结构差异。")
    add_para(doc, "小世界结构至少报告平均聚类系数 C、平均最短路径长度 L、模块度 Q、图稀疏度以及 small-worldness 指标 σ。随机图基线应重复采样并报告均值，避免单次随机参照造成波动。")
    add_table(
        doc,
        "表 9  可解释性与结构分析清单",
        ["分析对象", "输出形式", "需要回答的问题"],
        [
            ["动态神经图", "通道拓扑图、连接矩阵、时间变化曲线", "模型是否学习到稳定且具有任务区分度的连接模式"],
            ["认知单元激活", "类别分组热力图、轨迹图、聚类结果", "认知单元是否形成互补分工，是否具有可重复结构"],
            ["小世界结构", "C、L、Q、稀疏度、σ 及随机图对照", "结构约束是否带来局部聚类与全局整合之间的平衡"],
            ["不确定性样本", "高风险样本案例、置信度分布、决策类型", "模型是否能够识别易错样本并触发更保守动作"],
        ],
    )

    add_heading(doc, "5.8 Reproducibility Checklist", 2)
    add_para(doc, "为保证实验可以复核，最终版本需要同时保存数据划分文件、预处理配置、随机种子、模型超参数、最佳 checkpoint 选择规则、校准阈值选择规则、软件版本、硬件环境、参数量统计脚本和 batch size=1 的推理时延脚本。所有安全决策阈值必须仅在验证集上确定，测试集只用于一次性报告。")
    add_para(doc, "若实验采用 MOABB、Braindecode 或其他公开实现，应固定版本号，并记录对默认参数所做的修改。主文中保留核心设置，完整配置可放入附录或补充材料。")

    add_heading(doc, "6. Conclusion", 1)
    add_para(doc, "本文提出一种层级化代理脑框架，将 BCI 从脑信号分类扩展为神经—认知—意图—交互的动态建模问题。通过动态神经图、认知单元抽象、小世界结构约束、认知世界模型和不确定性感知决策机制，该框架为更准确、更可解释、更可信的脑机接口提供了统一建模思路。")
    add_para(doc, "需要进一步说明的是，当前实验部分应被定位为可复现实验方案和结果承载模板。只有在完成真实代码运行、统计检验和风险决策模拟后，才能将性能提升、安全收益和解释性结论写成实证结果。")

    add_heading(doc, "7. References", 1)
    refs = [
        "[1] Brunner C, Leeb R, Müller-Putz G R, Schlögl A, Pfurtscheller G. BCI Competition 2008 – Graz data set A. Graz University of Technology, 2008.",
        "[2] Goldberger A L, Amaral L A N, Glass L, et al. PhysioBank, PhysioToolkit, and PhysioNet: Components of a New Research Resource for Complex Physiologic Signals. Circulation, 2000, 101(23): e215–e220.",
        "[3] Schalk G, McFarland D J, Hinterberger T, Birbaumer N, Wolpaw J R. BCI2000: A General-Purpose Brain-Computer Interface System. IEEE Transactions on Biomedical Engineering, 2004, 51(6): 1034–1043.",
        "[4] Zheng W L, Lu B L. Investigating Critical Frequency Bands and Channels for EEG-Based Emotion Recognition with Deep Neural Networks. IEEE Transactions on Autonomous Mental Development, 2015, 7(3): 162–175.",
        "[5] Lawhern V J, Solon A J, Waytowich N R, et al. EEGNet: A Compact Convolutional Neural Network for EEG-Based Brain–Computer Interfaces. Journal of Neural Engineering, 2018, 15(5): 056013.",
        "[6] Schirrmeister R T, Springenberg J T, Fiederer L D J, et al. Deep Learning with Convolutional Neural Networks for EEG Decoding and Visualization. Human Brain Mapping, 2017, 38(11): 5391–5420.",
        "[7] Jayaram V, Barachant A. MOABB: Trustworthy Algorithm Benchmarking for BCIs. Journal of Neural Engineering, 2018, 15(6): 066011.",
        "[8] Carrara I, Papadopoulo T. Pseudo-online Framework for BCI Evaluation: A MOABB Perspective Using Various MI and SSVEP Datasets. Journal of Neural Engineering, 2024, 21(1): 016003.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.line_spacing = 1.0

    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    rebuild_doc()
    print(f"source={SOURCE_DOCX}")
    print(f"output={OUTPUT_DOCX}")
    print(f"figure={FIGURE_PATH}")
